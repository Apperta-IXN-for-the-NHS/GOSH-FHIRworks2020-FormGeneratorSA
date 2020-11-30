from docx import Document
from FHIRformGen import FHIRparser as fp

def createHeader(document, patientValues):
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    #the \t is to make it appear at the top right of document
    paragraph.text = "\t\tLast updated: " + patientValues.getLastUpdated()
    heading = document.add_heading('Medical record', 0)

def createIdentification(document, patientValues):
    document.add_heading("Identification", 1)
    p = document.add_paragraph("Name: " + patientValues.getGivenName() + "\n")
    p.add_run("Family Name: " + patientValues.getFamilyName() + "\n")
    p.add_run("patient ID: " + patientValues.getPatientID() + " \n")
    p.add_run("Medical record number: \n")
    p.add_run("Sex: " + patientValues.getSex()+ "\n")

def createDetails(document, patientValues):
    document.add_heading("Details", 1)
    p = document.add_paragraph("Birthday: " + patientValues.getBirthDate() + "\n")
    p.add_run("Language: " + + patientValues.getLanguage() + "\n")
    p.add_run("Birth Address: " + patientValues.getBirthAddress() + "\n")
    p.add_run("Marital Status: " + patientValues.getMaritalStatus() + "\n")

def createFormalDocs(document, patientValues):
    document.add_heading("Formal document numbers", 1)
    p = document.add_paragraph("Social Security: " + patientValues.getSocialSecurityNumber() +"\n")
    p.add_run("Driver License: " + patientValues.getDriversLicense() + "\n")
    p.add_run("Passport number: " + patientValues.getPassportNumber() +"\n")

def createContact(document, patientValues):
    document.add_heading("Contact", 1)
    p = document.add_paragraph("Number: " + patientValues.getContactDetails() +"\n")
    p.add_run("Address: " + patientValues.getCurrentAddress() +"\n")

def constructDocument(FHIRrecord):
    #calls the parser to get the information from the record
    patientValues = fp.FHIRParser(FHIRrecord)
    document = Document()
    createHeader(document, patientValues)
    createIdentification(document, patientValues)
    createFormalDocs(document, patientValues)
    createContact(document, patientValues)
    document.add_heading("Extra notes:", 1)
    return document

def createDocument(FHIRrecord, nameOfPDF):
    """ 
    function that acts as one of the entry points
    The function saves a docx based on the FHIR record provided

    parameters
    ----------
    FHIRrecord -> a JSON medical record obeying the FHIR standard
    nameOfPDF  -> string that the user wants the docx name to be

    """
    document = constructDocument(FHIRrecord)
    document.save(nameOfPDF + ".docx")

def getDocument(FHIRrecord):
    """ 
    function that acts as the second entry point
    The function returns a docx based on the FHIR record provided

    parameters
    ----------
    FHIRrecord -> a JSON medical record obeying the FHIR standard

    """
    return constructDocument(FHIRrecord)